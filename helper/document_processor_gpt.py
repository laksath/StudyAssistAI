import fitz  # PyMuPDF
import pdfplumber
from PIL import Image
import pytesseract
import io
import os
import magic
import docx
from datetime import datetime
from helper.gpt import completion
from concurrent.futures import ThreadPoolExecutor

def save_uploaded_file(uploaded_file):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename, file_extension = os.path.splitext(uploaded_file.name)
    unique_filename = f"{filename}_{timestamp}{file_extension}"
    filepath = os.path.join("temp_files", unique_filename)
    os.makedirs("temp_files", exist_ok=True)
    with open(filepath, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return filepath

def extract_text_images(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            images.append(image)
    return text, images

def extract_tables(pdf_path):
    tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables.extend(page.extract_tables())
    return tables

def ocr_images(images):
    def process_image(image):
        try:
            return pytesseract.image_to_string(image)
        except OSError:
            return ""

    with ThreadPoolExecutor() as executor:
        ocr_texts = list(executor.map(process_image, images))

    return ocr_texts

def clean_table_data(table):
    cleaned_table = []
    for row in table:
        cleaned_row = [str(cell) if cell is not None else '' for cell in row]
        cleaned_table.append(cleaned_row)
    return cleaned_table

def prepare_data_for_summary(text, tables, ocr_texts):
    combined_data = text
    for table in tables:
        cleaned_table = clean_table_data(table)
        combined_data += "\n" + "\n".join(["\t".join(row) for row in cleaned_table])
    combined_data += "\n" + "\n".join(ocr_texts)
    return combined_data

def extract_pdf_contents(pdf_path):
    text, images = extract_text_images(pdf_path)
    tables = extract_tables(pdf_path)
    ocr_texts = ocr_images(images)
    combined_data = prepare_data_for_summary(text, tables, ocr_texts)
    return combined_data

def extract_doc_text(doc):
    text = [paragraph.text for paragraph in doc.paragraphs]
    return "\n".join(text)

def extract_doc_tables(doc):
    tables = []
    for table in doc.tables:
        table_data = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(table_data)
    return tables

def extract_images(doc):
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image = rel.target_part.blob
            img = Image.open(io.BytesIO(image))
            images.append(img)
    return images
      
def extract_doc_contents(doc_path):
    doc = docx.Document(doc_path)
    text = extract_doc_text(doc)
    tables = extract_doc_tables(doc)
    images = extract_images(doc)
    ocr_texts = ocr_images(images)
    combined_data = prepare_data_for_summary(text, tables, ocr_texts)
    return combined_data

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)
    return text

def summary_prompt(text, task):
    if task == 'worksheet':
        return f"Summarize the following text in detail to prepare it for a large worksheet and return only the summarized text:\n\n{text}"
    elif task == 'mcq':
        return f"Summarize the following text in detail to prepare it for a large number of MCQs and return only the summarized text:\n\n{text}"
    elif task == 'comprehension':
        return f"Summarize the following text in detail to prepare it for a large number of comprehension-based questions and return only the summarized text:\n\n{text}"
    else:
        return f"Summarize the following text:\n\n{text}"

def generate_summary(api_key, text, task, max_tokens):
    prompt = summary_prompt(text, task)
    return completion(api_key, 'gpt-4o', prompt, max_tokens)

def get_file_type(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        return 'pdf'
    elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
        return 'image'
    elif file_extension in ['.doc', '.docx']:
        return 'word'
    
    mime = magic.Magic(mime=True)
    mime_type = mime.from_file(file_path)
    if mime_type == 'application/pdf':
        return 'pdf'
    elif mime_type.startswith('image/'):
        return 'image'
    elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        return 'word'
    else:
        return 'unknown'

def extract_summarized_document(file_path, api_key, task):
    if file_path =='':
        return ''

    file_type = get_file_type(file_path)
    if file_type == 'pdf':
      combined_data = extract_pdf_contents(file_path)
    elif file_type == 'image':
      combined_data = extract_text_from_image(file_path)
    elif file_type == 'word':
      combined_data = extract_doc_contents(file_path)
    
    summary = generate_summary(api_key, combined_data, task, None)
    
    return f'\n{summary}'